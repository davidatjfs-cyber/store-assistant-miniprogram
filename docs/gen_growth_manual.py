# -*- coding: utf-8 -*-
"""生成《HRMS + 小程序 增长模块使用手册》Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Document()

# ---------- 全局中文字体 ----------
def set_cn_font(run, name='微软雅黑', size=None, bold=None, color=None):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

# 设置正文默认字体
normal = DOC.styles['Normal']
normal.font.name = '微软雅黑'
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

ACCENT = (0x1F, 0x6F, 0xB2)
DARK = (0x22, 0x33, 0x44)

def h1(text):
    p = DOC.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_cn_font(run, size=18, bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    # 底部边框
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '1F6FB2')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = DOC.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size=14, bold=True, color=DARK)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def h3(text):
    p = DOC.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size=11.5, bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def para(text, bold=False, size=10.5, color=None):
    p = DOC.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(3)
    return p

def bullet(text, level=0):
    p = DOC.add_paragraph(style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_cn_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text):
    p = DOC.add_paragraph(style='List Number')
    run = p.add_run(text)
    set_cn_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def table(headers, rows, widths=None):
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext)
        set_cn_font(run, size=9.5, bold=True, color=(255, 255, 255))
        # 表头底色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F6FB2'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cn_font(run, size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def spacer():
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)

# ============================================================
# 封面
# ============================================================
title = DOC.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HRMS + 微信小程序\n增长模块使用手册')
set_cn_font(run, size=28, bold=True, color=ACCENT)
title.paragraph_format.space_before = Pt(120)
title.paragraph_format.space_after = Pt(20)

sub = DOC.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('门店会员增长 · 自动营销 · 私域触达 · 闭环对账')
set_cn_font(run, size=13, color=DARK)

meta = DOC.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(160)
run = meta.add_run('版本 v1.0　|　适用门店：马己仙（51866138）· 洪潮（64822111）\n编制日期：2026-06-04')
set_cn_font(run, size=10.5, color=DARK)

DOC.add_page_break()

# ============================================================
# 一、总体架构
# ============================================================
h1('一、总体架构与角色分工')

para('本系统由两套子系统协同构成，各司其职、双向打通：', bold=True)

table(
    ['子系统', '定位', '主要用户', '运行环境'],
    [
        ['微信小程序', '面向「顾客」的触点：扫码点餐、领券核销、会员手机号授权、卡包', '顾客 / 门店店员、店长', '腾讯云开发 CloudBase（云函数 + 云数据库）'],
        ['HRMS 增长后台', '面向「运营」的中枢：客群洞察、自动营销、内容创意、执行治理', '总部 / 管理员', '阿里云 ECS（Node 服务 nnyx.cc + Postgres）'],
    ],
    widths=[1.3, 3.0, 1.7, 1.6]
)

h2('1.1 数据如何打通')
para('两侧通过 HTTPS 接口双向同步，统一用密钥头 X-Miniprogram-Sync-Secret 鉴权：', size=10.5)
bullet('小程序 → HRMS：顾客在小程序内的高价值行为（手机号授权、支付成功、购券、核销）实时上报到 HRMS，用于会员画像与生命周期判定。')
bullet('HRMS → 小程序：运营在 HRMS 命中某会员后，可反向「指挥」小程序代发订阅消息、站内推券；并支持向沉睡客发送召回短信。')
bullet('HRMS → 小程序（拉取）：小程序每日从 HRMS 拉取 POS 真实消费金额，写回会员「消费金额」字段（客如云 POS 不直接推数据给小程序，靠此通道补齐）。')

h2('1.2 权限角色')
table(
    ['角色', '小程序', 'HRMS 增长后台'],
    [
        ['顾客', '点餐、领券、核销出示、查看卡包', '无'],
        ['店员 staff', '到店识别、补全客资、手动发券、扫码核销', '无'],
        ['店长 manager', '同店员 + 本店到店记录 / 看板', '无（数据在小程序端）'],
        ['管理员 admin', '创建营销规则、查看全量看板', '全部增长模块'],
    ],
    widths=[1.2, 3.2, 2.2]
)
para('说明：员工身份由小程序 staff 集合（按微信 openid 绑定、active=true）判定；未绑定者只能看到顾客界面。新增员工需运行 seedStaff 绑定后生效。', size=9.5, color=(0x88,0x55,0x00))

DOC.add_page_break()

# ============================================================
# 二、小程序增长模块
# ============================================================
h1('二、小程序端 增长模块')

para('小程序端的增长能力由「云函数」实现，分为 6 类。下文按"功能 / 触发方式 / 使用方法"说明。', size=10.5)

h2('2.1 会员识别与档案')
h3('① 手机号授权入会（saveUserPhone）')
bullet('功能：顾客点击授权手机号后，建立/更新会员档案 users；并检测是否为客如云导入的老会员（LegacyMembers），自动续接积分。')
bullet('触发：顾客在小程序点餐或领券流程中点「授权手机号」按钮。')
bullet('使用：顾客侧自动完成，无需运营操作。授权成功即上报 HRMS（事件 phone_authorized）。')

h3('② 到店识别（detectUserArrival）')
bullet('功能：顾客进入门店点餐页时，汇总其历史消费/到店/标签，生成画像并写入 user_arrival_logs；消费≥2 次的熟客会触发飞书提醒推送给本店员工。')
bullet('触发：顾客扫桌码进入点餐页自动触发（10 分钟内去重）。')
bullet('使用：店员可在小程序「最近到店」页查看（见 ③）；熟客到店时店员飞书群自动收到提醒。')

h3('③ 最近到店记录（getRecentArrivals）')
bullet('功能：按员工绑定门店展示最近到店顾客列表（仅店员/店长/管理员可见）。')
bullet('使用：店员在小程序员工界面打开「到店记录」，识别熟客、主动服务。')

h3('④ 补全客资（updateCustomerProfile）')
bullet('功能：店员为到店老客补「姓 + 性别」，系统派生中文称谓（如「张先生」），供短信/营销文案使用。')
bullet('使用：店员在顾客详情页填写姓氏与性别并保存。')

h3('⑤ 消费金额写回（syncPosConsumption）')
bullet('功能：每日从 HRMS 按手机号拉取 POS 真实消费聚合（总消费、订单数、近30天消费、最近到店、门店），写回会员档案的「消费金额」等字段。')
bullet('触发：每日 04:30 定时自动执行；也可手动执行，传 {dryRun:true} 仅预览不写库。')
bullet('使用：运营无需操作；如需立即刷新，可在云开发控制台手动调用该云函数。')

h2('2.2 客户管理（顾客列表）')
h3('客户列表（getCustomerList）')
bullet('功能：按门店展示会员列表，标注「生命周期阶段」标签，与 HRMS 口径 100% 一致：潜在新客 / 新客 / 活跃客 / 临界客 / 沉睡老客 / 流失客。')
bullet('口径：在该店「消费 / 到店 / 领券」任一发生即算该店客户；消费金额来自 POS 写回字段。')
bullet('使用：店长/管理员在小程序「客户」页查看，可按生命周期筛选并发起触达。')

h2('2.3 优惠券体系')
table(
    ['云函数', '功能', '使用者'],
    [
        ['createVoucherTemplate / getVoucherTemplates', '创建/查询券模板（面额、有效期、适用门店）', '管理员'],
        ['manualSendVoucher', '员工手动给指定顾客发券', '店员/店长/管理员'],
        ['getUserVouchers', '顾客查看「我的券」卡包（按状态/门店筛选）', '顾客'],
        ['verifyVoucher', '员工扫码核销券：校验门店/模板规则、防重复、记录核销日志', '店员/店长/管理员'],
        ['revertVoucher / reconcileRedemptions', '核销撤销与核销对账兜底', '管理员/系统'],
    ],
    widths=[2.4, 3.0, 1.4]
)
h3('核销操作流程')
numbered('顾客在「我的券」出示券二维码（voucher: 前缀）。')
numbered('店员用员工界面「扫码核销」扫描，系统校验门店与有效期。')
numbered('核销成功写 voucher_logs，并上报 HRMS（事件 coupon_redeemed），看板 ROI 实时更新。')

h2('2.4 自动营销引擎')
h3('营销规则（createMarketingRule / get / update / deleteMarketingRule）')
bullet('功能：管理员创建营销规则，定义触发条件与动作。')
bullet('触发类型：payment（支付后）/ inactivity（沉默N天）/ manual（手动单发）。')
bullet('动作类型：send_voucher（发券）。目标人群标签：潜在/新客/活跃/临界/沉睡/流失/VIP/常客/低价值。')
bullet('权限：仅管理员可创建/修改规则。')

h3('营销引擎执行（runMarketingEngine）')
bullet('post_payment 钩子：顾客支付成功后，匹配 payment 规则即时发券。')
bullet('inactivity_scan 钩子：每日全量扫描沉默会员，命中 inactivity 规则发券召回。')
bullet('manual 钩子：管理员对指定人群手动单发。')
bullet('daily_reconcile 钩子：每日补算 30 天指标与 ROI。')

h3('定时巡检（dailyCheckInactiveUsers）')
bullet('每日定时调用营销引擎：先跑沉默召回扫描，再跑 30天/ROI 兜底对账与系统监控。运营无需操作。')

h3('营销看板（getMarketingDashboard / getMarketingRules）')
bullet('功能：按规则汇总「发出 / 核销 / 营收 / ROI」，展示今日大盘与 TOP 规则。')
bullet('使用：管理员在小程序看板页查看各规则效果，判断是否调整或下线。')

h2('2.5 主动触达通道')
table(
    ['云函数', '渠道', '说明'],
    [
        ['sendSubscribeMessage', '微信订阅消息', '发「领券通知 / 到期提醒」；仅能发给点过订阅授权且有剩余次数的用户'],
        ['growthSubscribePush', 'HRMS→订阅消息网关', 'HRMS 解析 openid 后委托小程序代发订阅消息'],
        ['growthMemberCoupon', 'HRMS→站内推券网关', 'HRMS 命中会员后直接发券进卡包，带幂等键防重复'],
        ['sendWinbackCampaign', '阿里云短信', '为沉睡客生成带 6 位短码的现金券并由 HRMS 发召回短信'],
        ['notifyRegularCustomerArrival', '飞书提醒', '熟客到店推送给本店员工飞书群'],
    ],
    widths=[2.2, 1.7, 3.0]
)

h2('2.6 可靠性保障（事件同步加固）')
para('为确保小程序高价值事件不丢失，系统采用「实时重试 + 兜底对账」双保险：', bold=True)
bullet('实时重试：事件上报 HRMS 失败时自动重试 3 次（指数退避）；4xx 鉴权/参数类错误不重试。')
bullet('落库兜底：重试仍失败的事件写入 hrms_event_outbox 队列，绝不静默丢弃。')
bullet('定时重投（reconcileHrmsEvents）：每 10 分钟扫描队列重新投递；HRMS 按幂等键去重，重投不会重复计数；超 8 次或永久错误标记 failed 待人工排查。')
para('运营无需关注此机制，属系统级自愈能力。', size=9.5, color=(0x55,0x55,0x55))

DOC.add_page_break()

# ============================================================
# 三、HRMS 增长后台
# ============================================================
h1('三、HRMS 增长后台 使用指南')

para('入口：HRMS 系统左侧菜单「📈 增长看板」（仅管理员 / 总部可见）。页面顶部有 6 个标签页。', size=10.5)

table(
    ['标签页', '子模块', '核心用途'],
    [
        ['看板', '看板', 'POS 与小程序双源数据总览、漏斗、预警、最佳触达时段'],
        ['客群洞察', '用户画像 / POS消费 / 企微客户', '看清「客是谁、消费力如何、私域绑定情况」'],
        ['自动营销', '规则引擎 / A/B测试', '配置自动触达规则、对照实验'],
        ['内容中心', '活动管理 / 海报创意 / 公域品宣 / 内容系统', '营销活动与创意素材生产'],
        ['执行中心', 'AI建议 / 执行记录', '审批并执行 AI 营销建议、查执行流水'],
        ['设置治理', '企微配置 / 营销约束', '门店企微对接、营销频次与额度护栏'],
    ],
    widths=[1.2, 2.6, 3.0]
)

h2('3.1 看板')
bullet('POS 系统数据卡：来自客如云 POS 的真实消费（订单数、营收、客单价、堂食率）。')
bullet('小程序数据卡：会员数、券发放/核销、活动漏斗等小程序侧指标。')
bullet('流失预警：临界/沉睡会员预警，可一键转入营销。')
bullet('天气/时令、最佳触达时段、复购触发：辅助决定「何时、向谁、推什么」。')
bullet('行动建议：系统聚合的 AI 营销建议入口（在执行中心审批执行）。')
bullet('使用：进入页面先选「门店 / 活动 / 时间范围」筛选器，再读各卡片。')

h2('3.2 客群洞察')
h3('用户画像（profiles）')
bullet('功能：基于 POS 消费 + 小程序行为计算的会员画像与生命周期分群，可触发 recompute 重算。')
h3('POS 消费（pos）')
bullet('功能：展示按手机号聚合的 POS 消费明细，是小程序「消费金额」的数据源。')
h3('企微客户（wecom）')
bullet('功能：企微好友客户管理，支持 CSV/Excel 导入、飞书多维表格自动同步、手动添加；展示绑定/未绑定统计。')
bullet('使用：刷新看板查看统计 → 对未绑定客户「推入会」→ 在「设置治理-企微配置」维护门店密钥。')

h2('3.3 自动营销')
h3('规则引擎（automarketing）')
bullet('功能：配置 HRMS 侧触达规则（人群条件 + 动作 + 频次约束），命中后通过订阅消息/站内推券/短信触达会员。')
bullet('使用：新建规则→设定目标人群与触达内容→预览覆盖会员数→审批生效。')
h3('A/B 测试（abtests）')
bullet('功能：对同一目标人群拆分对照组，比较不同文案/券面额的转化与 ROI。')

h2('3.4 内容中心')
bullet('活动管理（campaigns）：创建营销活动、查看活动漏斗（曝光→领取→核销→营收）的真实回流归因。')
bullet('海报创意（posters）：海报模板与生成海报管理，供门店私域/公域投放。')
bullet('公域品宣（public）：公域渠道与品宣任务管理（如点评、社媒）。')
bullet('内容系统（contentsys）：内容素材库与内容效果统计。')

h2('3.5 执行中心')
h3('AI 建议（actions）')
bullet('功能：系统/AI 产出的营销动作建议，管理员可「执行 / 编辑后执行 / 忽略」。')
bullet('使用：逐条审阅建议→点「执行」即按建议触达，或「编辑后执行」微调文案/人群后再发。')
h3('执行记录（exec-logs）')
bullet('功能：所有已执行营销动作的流水与结果，用于追溯与复盘。')

h2('3.6 设置治理')
h3('企微配置（wecomconfig）')
bullet('功能：维护各门店企业微信对接配置、同步企微客户。')
bullet('安全：企微 corp_secret 等密钥存于数据库，切勿硬编码或外泄。')
h3('营销约束（constraints）')
bullet('功能：设置营销频次上限、单客触达额度、券预算等护栏，防止过度打扰与超支。')
bullet('使用：先在此设好约束，自动营销与 AI 建议执行时会自动遵守。')

h2('3.7 沉睡客召回短信（HRMS→小程序）')
bullet('入口：HRMS 接口 /api/growth/winback/send-sms（由召回流程调用）。')
bullet('流程：HRMS 选定沉睡会员名单 → 小程序为每人生成带 6 位短码的无门槛现金券 → HRMS 用阿里云短信发出 → 顾客凭短码到店核销 → 回流归因到对应活动。')
bullet('注意：测试短信仅可发本人手机号，避免打扰真实顾客。')

DOC.add_page_break()

# ============================================================
# 四、典型业务流程
# ============================================================
h1('四、典型业务场景操作流程')

h2('场景 A：新客首次到店转化')
numbered('顾客扫桌码点餐 → 授权手机号入会（saveUserPhone，建档并上报 HRMS）。')
numbered('支付成功 → 营销引擎 payment 规则即时发「新客券」进卡包（runMarketingEngine）。')
numbered('订阅消息提醒顾客「领券成功」（sendSubscribeMessage）。')
numbered('顾客下次到店出示券 → 店员扫码核销（verifyVoucher）→ 看板 ROI 更新。')

h2('场景 B：沉睡老客召回')
numbered('HRMS 看板「流失预警」识别沉睡会员，或自动营销 inactivity 规则每日扫描命中。')
numbered('对名单发起召回：站内推券（growthMemberCoupon）或召回短信（sendWinbackCampaign）。')
numbered('顾客凭券/短码到店核销 → 回流归因到召回活动 → 在活动漏斗看转化。')

h2('场景 C：熟客到店服务')
numbered('熟客（消费≥2次）到店 → detectUserArrival 触发飞书提醒到店员群。')
numbered('店员在「最近到店」查看画像，必要时补全客资（updateCustomerProfile）。')
numbered('店员可手动发券（manualSendVoucher）做即时关怀。')

h2('场景 D：会员消费金额对账')
numbered('客如云 POS 订单数据进入 HRMS 的 pos_orders。')
numbered('每日 04:30，小程序 syncPosConsumption 按手机号拉取并写回会员「消费金额」。')
numbered('小程序客户列表与 HRMS 看板口径一致，生命周期分群随之更新。')

DOC.add_page_break()

# ============================================================
# 五、常见问题
# ============================================================
h1('五、常见问题与维护')

qa = [
    ('员工看到的是顾客界面，看不到员工功能？',
     '该员工的 staff 记录未绑定微信 openid。请确认其已用本人微信授权手机号后，运行 seedStaff 绑定，刷新即生效。'),
    ('小程序里会员「消费金额」为 0 或偏低？',
     '消费金额来自 POS 写回。请确认客如云 POS 数据已进入 HRMS、且会员手机号与 POS 单据一致；可手动执行 syncPosConsumption 立即刷新。'),
    ('订阅消息发不出去？',
     '微信限制：只能发给点过订阅授权且仍有剩余次数的用户。返回 43101 属未授权，不算系统故障，需引导顾客再次授权。'),
    ('高价值事件会不会丢？会不会重复计数？',
     '不会。失败事件落 hrms_event_outbox 由 reconcileHrmsEvents 每 10 分钟重投；HRMS 按幂等键去重，重投不重复计数。'),
    ('POS 数据更新不及时有什么影响？',
     '会导致会员生命周期阶段被错判，进而触达发错人、发无效券，ROI 统计失真、误判活动好坏。应加强 POS 数据时效性。'),
    ('谁能进入 HRMS 增长后台？',
     '仅管理员 / 总部角色可见「增长看板」菜单，店长/店员的客户与到店数据在小程序端查看。'),
]
for q, a in qa:
    h3('Q：' + q)
    para('A：' + a, size=10.5)

spacer()
h2('附：关键定时任务')
table(
    ['任务', '频率', '作用'],
    [
        ['syncPosConsumption', '每日 04:30', '写回会员 POS 消费金额'],
        ['dailyCheckInactiveUsers', '每日', '沉默召回扫描 + 30天/ROI 对账'],
        ['reconcileHrmsEvents', '每 10 分钟', '失败事件兜底重投'],
    ],
    widths=[2.4, 1.6, 3.0]
)

# 页脚
section = DOC.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
frun = fp.add_run('HRMS + 小程序 增长模块使用手册 · v1.0 · 内部资料')
set_cn_font(frun, size=8, color=(0x99, 0x99, 0x99))

OUT = '/Users/xieding/store-assistant-miniprogram/docs/增长模块使用手册.docx'
DOC.save(OUT)
print('saved:', OUT)
